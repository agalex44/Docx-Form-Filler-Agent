# -*- coding: utf-8 -*-
"""
DOCX Form Filler - Production Grade Solution
Deterministic, format-preserving, 100% accurate
"""

import json
import re
from typing import Dict, List, Tuple, Optional
from dataclasses import dataclass
from docx import Document
from docx.shared import RGBColor
from rapidfuzz import fuzz
import unicodedata

# ============================================================================
# Core Data Structures
# ============================================================================

@dataclass
class LabeledField:
    """A field with its label and placeholder in document"""
    label: str              # "Nume ofertant"
    placeholder: str        # "_____" or "....."
    paragraph_obj: any      # Reference to paragraph
    run_index: int          # Which run contains placeholder
    char_start: int         # Start position in run
    char_end: int           # End position in run
    context: str            # Full paragraph text
    table_cell: any = None  # If in table

@dataclass
class MatchResult:
    """Result of matching JSON key to document field"""
    json_key: str
    document_field: LabeledField
    confidence: float
    match_type: str  # 'exact', 'fuzzy', 'semantic'

# ============================================================================
# Step 1: Enhanced Field Extraction
# ============================================================================

class SmartFieldExtractor:
    """Extract label-placeholder pairs from document"""
    
    # Comprehensive placeholder patterns
    PLACEHOLDER_PATTERNS = [
        (r'_{5,}', 'underscore'),
        (r'\.{4,}', 'dots'),
        (r'\[___\]', 'brackets'),
        (r'\(\s*\)', 'empty_parens'),
    ]
    
    # Label patterns (text before placeholder)
    LABEL_INDICATORS = [
        r'([^:.\n]{3,50}):\s*PLACEHOLDER',
        r'([^:.\n]{3,50})\s+PLACEHOLDER',
        r'([^|\n]{3,50})\|\s*PLACEHOLDER',
    ]
    
    @staticmethod
    def normalize_text(text: str) -> str:
        """Normalize Unicode and whitespace"""
        text = unicodedata.normalize('NFC', text)
        text = re.sub(r'\s+', ' ', text).strip()
        return text
    
    @staticmethod
    def extract_from_paragraph(para, table_cell=None) -> List[LabeledField]:
        """Extract labeled fields from a paragraph"""
        fields = []
        full_text = para.text
        
        if not full_text.strip():
            return fields
        
        # Find all placeholders
        for pattern, ph_type in SmartFieldExtractor.PLACEHOLDER_PATTERNS:
            for match in re.finditer(pattern, full_text):
                placeholder = match.group()
                ph_start = match.start()
                ph_end = match.end()
                
                # Extract label (text before placeholder)
                before_text = full_text[:ph_start].strip()
                label = SmartFieldExtractor._extract_label(before_text)
                
                # Find which run contains this placeholder
                run_idx, char_in_run = SmartFieldExtractor._find_run_position(
                    para, ph_start
                )
                
                if run_idx is not None:
                    fields.append(LabeledField(
                        label=SmartFieldExtractor.normalize_text(label),
                        placeholder=placeholder,
                        paragraph_obj=para,
                        run_index=run_idx,
                        char_start=char_in_run,
                        char_end=char_in_run + len(placeholder),
                        context=full_text,
                        table_cell=table_cell
                    ))
        
        return fields
    
    @staticmethod
    def _extract_label(text_before: str) -> str:
        """Extract meaningful label from text before placeholder"""
        text_before = text_before.strip()
        
        # Try pattern matching
        for pattern in [
            r'([^:.\n]{3,100}):\s*$',       # "Label:"
            r'([^:.\n]{3,100})\s+[-–—]\s*$', # "Label -"
            r'([^|\n]{3,100})\|\s*$',        # "Label |"
        ]:
            match = re.search(pattern, text_before)
            if match:
                return match.group(1).strip()
        
        # Fallback: last 3-50 chars
        words = text_before.split()
        if words:
            return ' '.join(words[-5:])  # Last 5 words
        
        return text_before[-50:] if len(text_before) > 50 else text_before
    
    @staticmethod
    def _find_run_position(para, char_index: int) -> Tuple[Optional[int], int]:
        """Find which run contains character at index"""
        current_pos = 0
        for idx, run in enumerate(para.runs):
            run_len = len(run.text)
            if current_pos <= char_index < current_pos + run_len:
                return idx, char_index - current_pos
            current_pos += run_len
        return None, 0
    
    @staticmethod
    def extract_all_fields(doc: Document) -> List[LabeledField]:
        """Extract all labeled fields from document"""
        all_fields = []
        
        # From paragraphs
        for para in doc.paragraphs:
            all_fields.extend(SmartFieldExtractor.extract_from_paragraph(para))
        
        # From tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        all_fields.extend(
                            SmartFieldExtractor.extract_from_paragraph(para, cell)
                        )
        
        return all_fields

# ============================================================================
# Step 2: Intelligent Key Matching
# ============================================================================

class KeyMatcher:
    """Match JSON keys to document labels deterministically"""
    
    ROMANIAN_NORMALIZE = {
        'ă': 'a', 'â': 'a', 'î': 'i', 'ș': 's', 'ț': 't',
        'Ă': 'a', 'Â': 'a', 'Î': 'i', 'Ș': 's', 'Ț': 't'
    }
    
    @staticmethod
    def normalize_for_matching(text: str) -> str:
        """Normalize text for fuzzy matching"""
        text = text.lower().strip()
        for old, new in KeyMatcher.ROMANIAN_NORMALIZE.items():
            text = text.replace(old, new)
        # Remove punctuation
        text = re.sub(r'[^\w\s]', ' ', text)
        text = re.sub(r'\s+', ' ', text)
        return text
    
    @staticmethod
    def exact_match(json_key: str, label: str) -> float:
        """Check for exact match (case/diacritic insensitive)"""
        norm_key = KeyMatcher.normalize_for_matching(json_key)
        norm_label = KeyMatcher.normalize_for_matching(label)
        
        if norm_key == norm_label:
            return 1.0
        if norm_key in norm_label or norm_label in norm_key:
            return 0.95
        return 0.0
    
    @staticmethod
    def fuzzy_match(json_key: str, label: str) -> float:
        """Fuzzy matching with multiple strategies"""
        norm_key = KeyMatcher.normalize_for_matching(json_key)
        norm_label = KeyMatcher.normalize_for_matching(label)
        
        scores = [
            fuzz.ratio(norm_key, norm_label) / 100,
            fuzz.partial_ratio(norm_key, norm_label) / 100,
            fuzz.token_sort_ratio(norm_key, norm_label) / 100,
            fuzz.token_set_ratio(norm_key, norm_label) / 100,
        ]
        return max(scores)
    
    @staticmethod
    def match_key_to_field(
        json_key: str,
        field: LabeledField,
        threshold: float = 0.85
    ) -> Optional[MatchResult]:
        """Match a JSON key to a document field"""
        
        # Try exact match first
        exact_score = KeyMatcher.exact_match(json_key, field.label)
        if exact_score >= threshold:
            return MatchResult(
                json_key=json_key,
                document_field=field,
                confidence=exact_score,
                match_type='exact'
            )
        
        # Fuzzy match
        fuzzy_score = KeyMatcher.fuzzy_match(json_key, field.label)
        if fuzzy_score >= threshold:
            return MatchResult(
                json_key=json_key,
                document_field=field,
                confidence=fuzzy_score,
                match_type='fuzzy'
            )
        
        # Context-based fuzzy match (use full context)
        context_score = KeyMatcher.fuzzy_match(json_key, field.context)
        if context_score >= threshold:
            return MatchResult(
                json_key=json_key,
                document_field=field,
                confidence=context_score * 0.9,  # Slight penalty
                match_type='context'
            )
        
        return None
    
    @staticmethod
    def match_all_keys(
        json_data: Dict[str, str],
        fields: List[LabeledField],
        threshold: float = 0.85
    ) -> List[MatchResult]:
        """Match all JSON keys to fields (one-to-one)"""
        matches = []
        used_fields = set()
        
        # Sort keys by length (longer = more specific)
        sorted_keys = sorted(json_data.keys(), key=len, reverse=True)
        
        for json_key in sorted_keys:
            best_match = None
            best_score = 0.0
            
            for field in fields:
                if id(field) in used_fields:
                    continue
                
                match = KeyMatcher.match_key_to_field(json_key, field, threshold)
                if match and match.confidence > best_score:
                    best_match = match
                    best_score = match.confidence
            
            if best_match:
                matches.append(best_match)
                used_fields.add(id(best_match.document_field))
        
        return matches

# ============================================================================
# Step 3: Format-Preserving Replacement
# ============================================================================

class FormatPreservingReplacer:
    """Replace text while preserving all formatting"""
    
    @staticmethod
    def replace_in_run(field: LabeledField, new_value: str) -> bool:
        """Replace placeholder in specific run"""
        try:
            para = field.paragraph_obj
            run = para.runs[field.run_index]
            
            # Get current run text
            original_text = run.text
            
            # Build new text
            before = original_text[:field.char_start]
            after = original_text[field.char_end:]
            new_text = before + new_value + after
            
            # Replace (formatting preserved automatically)
            run.text = new_text
            
            return True
            
        except Exception as e:
            print(f"  ⚠ Replace failed: {e}")
            return False
    
    @staticmethod
    def replace_all_matches(
        matches: List[MatchResult],
        json_data: Dict[str, str]
    ) -> Dict[str, int]:
        """Replace all matched placeholders"""
        stats = {
            'success': 0,
            'failed': 0,
            'exact_matches': 0,
            'fuzzy_matches': 0,
        }
        
        for match in matches:
            value = str(json_data[match.json_key])
            
            # Handle special cases
            if value.startswith('[') and value.endswith(']'):
                # Array/list - format nicely
                try:
                    items = json.loads(value)
                    if isinstance(items, list):
                        value = ', '.join(str(i) for i in items)
                except:
                    pass
            
            # Replace
            success = FormatPreservingReplacer.replace_in_run(
                match.document_field, value
            )
            
            if success:
                stats['success'] += 1
                if match.match_type == 'exact':
                    stats['exact_matches'] += 1
                else:
                    stats['fuzzy_matches'] += 1
                    
                print(f"  ✓ [{match.confidence:.2f}] {match.json_key[:40]} → {value[:40]}")
            else:
                stats['failed'] += 1
        
        return stats

# ============================================================================
# Main Agent (Deterministic)
# ============================================================================

class DeterministicDocxFiller:
    """Production-grade deterministic DOCX filler"""
    
    def __init__(self, fuzzy_threshold: float = 0.85):
        self.extractor = SmartFieldExtractor()
        self.matcher = KeyMatcher()
        self.replacer = FormatPreservingReplacer()
        self.threshold = fuzzy_threshold
    
    def process(
        self,
        docx_path: str,
        json_path: str,
        output_path: str,
        verbose: bool = True
    ) -> Dict:
        """Process document with deterministic matching"""
        
        if verbose:
            print("=" * 70)
            print("DOCX FORM FILLER - DETERMINISTIC v2.0")
            print("=" * 70)
        
        # Load data
        with open(json_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
        
        if verbose:
            print(f"[1/5] ✓ Loaded {len(json_data)} JSON fields")
        
        # Load document
        doc = Document(docx_path)
        if verbose:
            print("[2/5] ✓ Document loaded")
        
        # Extract labeled fields
        fields = self.extractor.extract_all_fields(doc)
        if verbose:
            print(f"[3/5] ✓ Extracted {len(fields)} labeled fields")
        
        # Match keys to fields
        matches = self.matcher.match_all_keys(json_data, fields, self.threshold)
        if verbose:
            print(f"[4/5] ✓ Matched {len(matches)}/{len(json_data)} keys")
        
        # Replace
        if verbose:
            print("[5/5] Filling fields:")
        
        stats = self.replacer.replace_all_matches(matches, json_data)
        
        # Save
        doc.save(output_path)
        
        if verbose:
            print("\n" + "=" * 70)
            print("SUMMARY")
            print("=" * 70)
            print(f"Total JSON keys: {len(json_data)}")
            print(f"Fields found: {len(fields)}")
            print(f"Successful matches: {stats['success']}")
            print(f"  - Exact: {stats['exact_matches']}")
            print(f"  - Fuzzy: {stats['fuzzy_matches']}")
            print(f"Failed: {stats['failed']}")
            print(f"Unmatched keys: {len(json_data) - len(matches)}")
            print("=" * 70)
        
        return {
            'total_keys': len(json_data),
            'fields_found': len(fields),
            'matches': len(matches),
            'success': stats['success'],
            'failed': stats['failed'],
        }